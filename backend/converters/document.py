import os
from docx import Document
from fpdf import FPDF
from pypdf import PdfReader

TEXT_INPUTS = {".txt", ".md", ".html", ".csv", ".rtf"}
PDF_INPUTS = {".pdf"}

SUPPORTED_OUTPUTS = {
    ".txt", ".md", ".html", ".rtf", ".csv", ".docx", ".pdf"
}


def convert_document(input_path, output_path):
    input_ext = os.path.splitext(input_path)[1].lower()
    output_ext = os.path.splitext(output_path)[1].lower()

    if output_ext not in SUPPORTED_OUTPUTS:
        raise ValueError(f"Unsupported output format: {output_ext}")

    text = _extract_text(input_path, input_ext)

    if output_ext in {".txt", ".md"}:
        _write_text(output_path, text)

    elif output_ext == ".html":
        _write_html(output_path, text)

    elif output_ext == ".rtf":
        _write_rtf(output_path, text)

    elif output_ext == ".csv":
        _write_csv(output_path, text)

    elif output_ext == ".docx":
        _write_docx(output_path, text)

    elif output_ext == ".pdf":
        _write_pdf(output_path, text)

    return output_path

def _extract_text(path, ext):
    if ext in TEXT_INPUTS:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    if ext == ".pdf":
        reader = PdfReader(path)
        return "\n".join(
            page.extract_text() or ""
            for page in reader.pages
        )

    if ext == ".docx":
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    raise ValueError(f"Unsupported input format: {ext}")

def _write_text(path, text):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)

def _write_plain(path, text):
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)


def _write_html(path, text):
    html = f"""<!DOCTYPE html>
<html>
<head><meta charset="utf-8"></head>
<body><pre>{text}</pre></body>
</html>"""
    _write_plain(path, html)


def _write_rtf(path, text):
    rtf = (
        r"{\rtf1\ansi "
        + text.replace("\\", r"\\")
              .replace("{", r"\{")
              .replace("}", r"\}")
              .replace("\n", r"\line ")
        + "}"
    )
    _write_plain(path, rtf)


def _write_csv(path, text):
    with open(path, "w", encoding="utf-8") as f:
        for line in text.splitlines():
            f.write(",".join(line.split()) + "\n")


def _write_docx(path, text):
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(path)


def _write_pdf(path, text):
    pdf = FPDF()
    pdf.set_auto_page_break(True, 15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, text)
    pdf.output(path)
